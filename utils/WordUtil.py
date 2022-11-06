import logging
import re
from docx import Document


class WordUtil:

    @staticmethod
    def docx_find_section(doc_obj: Document):
        for p in doc_obj.paragraphs:
            logging.info(f'style->{p.style.name}, ptxt->{p.text}')

    @staticmethod
    def docx_replace_regex(doc_obj: Document, regex, replace):
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    WordUtil.docx_replace_regex(cell, regex, replace)

    @staticmethod
    def docx_replace_dic(doc: Document, dic):
        for word, replacement in dic.items():
            word_re = re.compile(word)
            WordUtil.docx_replace_regex(doc, word_re, replacement)

    @staticmethod
    def docx_find_runs(doc: Document, word, follow=1):
        result = []
        regex = re.compile(word)
        for p in doc.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                total_runs = range(len(inline))
                for i in total_runs:
                    if regex.search(inline[i].text):
                        if i + follow in total_runs:
                            result.append(
                                inline[i].text + inline[i + follow].text)
                        else:
                            result.append(inline[i].text)
        return result
